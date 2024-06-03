from docx import Document
from docx.shared import Inches
import pyttsx3
import datetime

now = datetime.datetime.now()

def speak(text):
    pyttsx3.speak(text)

document = Document()

# Profile Pictures
document.add_picture("future_profile1.jpg", width=Inches(1.5))

# Name, Phone Numbers, and Email inputs
name = input("What is your name: ")
speak(f"Hello, {name}! I hope you are doing well. Please enter your phone number?")
phone_number = input("What is your phone number: ")
email = input("What is your email: ")

document.add_paragraph(
    name + " | " + phone_number + " | " + email)

# About me Section
document.add_heading("About Me")
about_me = input("Tell me about yourself: ")
document.add_paragraph(about_me)

# Work Experience Section
document.add_heading("Work Experience")
p = document.add_paragraph()

company = input("Enter Company Name: ")
from_date = input("Enter your start date: ")
to_date = input("Enter your end date: ")

p.add_run(company + " ").bold = True
p.add_run(from_date + "-" + to_date + "\n").italic = True

experience_details = input(f"Enter your experience at {company}: ")
p.add_run(experience_details)

# More Experiences While Loop
while True:
    has_more_experiences = input("Do you have more work experiences? (Yes/No): ")
    if has_more_experiences.lower() == "yes":
        p = document.add_paragraph()

        company = input("Enter Company Name: ")
        from_date = input("Enter your start date: ")
        to_date = input("Enter your end date: ")

        p.add_run(company + " ").bold = True
        p.add_run(from_date + "-" + to_date + "\n").italic = True

        experience_details = input(f"Enter your experience at {company}: ")
        p.add_run(experience_details)
    else:
        break

# List of Technical Skills
document.add_heading("Technical Skills")
tskills = input("Enter A Technical Skill: ")
p = document.add_paragraph(tskills)
p.style = "List Bullet"

# List of Technical Skills While Loop
while True:
    has_more_techskills = input("Do you have more Technical Skills? (Yes/No): ")
    if has_more_techskills.lower() == "yes":
        tskills = input("Enter A Technical Skill: ")
        p = document.add_paragraph(tskills)
        p.style = "List Bullet"
    else:
        break

# List of Soft Skills
document.add_heading("Soft Skills")
sskills = input("Enter A Soft Skill: ")
p = document.add_paragraph(sskills)
p.style = "List Bullet"

# List of Technical Skills While Loop
while True:
    has_more_softskills = input("Do you have more Soft Skills? (Yes/No): ")
    if has_more_softskills.lower() == "yes":
        sskills = input("Enter A Soft Skill: ")
        p = document.add_paragraph(sskills)
        p.style = "List Bullet"
    else:
        break

# Footer
section = document.sections[0]
footer = section.footer
p = footer.paragraphs[0]
p.text = f"CV generated for {name} on {now.strftime("%Y-%m-%d %H:%M:%S")}"

document.save("cv.docx")